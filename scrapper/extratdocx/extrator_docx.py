import docx
import pandas as pd
import os
import re
from datetime import datetime

def extrair_texto_e_tabelas(caminho_docx):
    """
    Extrai texto e tabelas de um arquivo DOCX.
    Retorna uma lista com todas as strings encontradas.
    """
    try:
        doc = docx.Document(caminho_docx)

        # Extrai texto de parágrafos
        textos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Extrai texto das tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [celula.text.strip() for celula in linha.cells if celula.text.strip()]
                if celulas:
                    textos.append(" | ".join(celulas))

        return textos

    except Exception as e:
        print(f" Erro ao processar '{caminho_docx}': {e}")
        return []


def extrair_numeros(textos):
    """
    Extrai números (inteiros e decimais) de uma lista de strings.
    """
    numeros = []
    for linha in textos:
        encontrados = re.findall(r"[-+]?\d*\.\d+|\d+", linha)
        numeros.extend([float(n) for n in encontrados])
    return numeros


def processar_docx_pasta(pasta="."):
    """
    Lê todos os arquivos .docx da pasta, extrai textos, números,
    calcula estatísticas e salva em CSV.
    """
    arquivos = [f for f in os.listdir(pasta) if f.lower().endswith(".docx")]
    if not arquivos:
        print("Nenhum arquivo .docx encontrado na pasta.")
        return

    todos_resultados = []

    for arquivo in arquivos:
        caminho = os.path.join(pasta, arquivo)
        print(f"\nProcessando: {arquivo}")

        textos = extrair_texto_e_tabelas(caminho)
        numeros = extrair_numeros(textos)

        # Salva texto extraído em CSV individual
        nome_csv_texto = f"{os.path.splitext(arquivo)[0]}_texto_extraido.csv"
        df_texto = pd.DataFrame({"Texto": textos})
        df_texto.to_csv(nome_csv_texto, index=False, encoding="utf-8-sig")
        print(f"Texto salvo em: {nome_csv_texto}")

        # Se tiver números, calcula estatísticas
        if numeros:
            df_num = pd.DataFrame(numeros, columns=["Valores"])
            estatisticas = df_num.describe().transpose()
            estatisticas["arquivo"] = arquivo
            todos_resultados.append(estatisticas)

            nome_csv_est = f"{os.path.splitext(arquivo)[0]}_estatisticas.csv"
            estatisticas.to_csv(nome_csv_est, index=False, encoding="utf-8-sig")
            print(f"Estatísticas salvas em: {nome_csv_est}")
        else:
            print("Nenhum número encontrado neste DOCX.")

    # Combina todas as estatísticas em um CSV geral (se houver)
    if todos_resultados:
        df_final = pd.concat(todos_resultados, ignore_index=True)
        nome_csv_geral = f"estatisticas_docx_geral_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        df_final.to_csv(nome_csv_geral, index=False, encoding="utf-8-sig")
        print(f"\n Estatísticas combinadas salvas em: {nome_csv_geral}")
    else:
        print("\n Nenhum número encontrado em nenhum documento.")


if __name__ == "__main__":
    processar_docx_pasta(".")
