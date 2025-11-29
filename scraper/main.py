"""
main.py
Dependências:
    pip install pandas python-docx pdfplumber tabulate openpyxl tqdm streamlit
"""

from __future__ import annotations
import argparse
import logging
import os
from pathlib import Path
import re
from typing import List, Tuple, Optional, Dict, Any
from datetime import datetime

import pandas as pd

#tentativas opcionais de importação
try:
    import docx
except:
    docx = None

try:
    import pdfplumber
except:
    pdfplumber = None

from tqdm import tqdm
from tabulate import tabulate


#logging
LOG_FMT = "%(asctime)s %(levelname)-8s %(message)s"
logging.basicConfig(level=logging.INFO, format=LOG_FMT)
logger = logging.getLogger("main_processor")

#funções utilitárias
def safe_mkdir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)

def copy_csvs_to_data(out_dir: Path, data_dir: Path) -> None:
    """Copia todos os CSVs gerados para a pasta data/."""
    safe_mkdir(data_dir)

    for csv_file in out_dir.rglob("*.csv"):
        dest = data_dir / csv_file.name
        try:
            import shutil
            shutil.copy(csv_file, dest)
        except Exception as e:
            logger.error(f"Falha ao copiar {csv_file} para {dest}: {e}")


def is_number_series(s: pd.Series) -> bool:
    return pd.api.types.is_numeric_dtype(s)


def compute_descriptive_stats(df: pd.DataFrame) -> pd.DataFrame:
    """Calcula estatísticas robustas (describe + CV)."""
    if df.empty:
        return pd.DataFrame()

    desc = df.describe(include='all').transpose()
    numeric_cols = df.select_dtypes(include='number').columns

    for col in numeric_cols:
        mean = df[col].mean()
        std = df[col].std()
        cv = std / mean if (pd.notna(mean) and mean != 0) else pd.NA
        desc.loc[col, "cv"] = cv

    # arredonda floats
    float_cols = desc.select_dtypes(include=['float']).columns
    for c in float_cols:
        desc[c] = desc[c].round(6)

    return desc


def textual_summary(df: pd.DataFrame) -> pd.DataFrame:
    """Resumo textual interpretável por coluna."""
    rows = []
    for col in df.columns:
        s = df[col]
        if is_number_series(s):
            mean = s.mean()
            std = s.std()
            cv = std / mean if (pd.notna(mean) and mean != 0) else pd.NA
            rows.append({
                "coluna": col,
                "tipo": str(s.dtype),
                "media": mean,
                "std": std,
                "cv": cv
            })
        else:
            moda = s.mode().iat[0] if not s.mode().empty else pd.NA
            rows.append({
                "coluna": col,
                "tipo": str(s.dtype),
                "unicos": s.nunique(dropna=True),
                "top": moda,
                "freq": s.value_counts().iloc[0] if not s.value_counts().empty else pd.NA
            })

    return pd.DataFrame(rows)


#CSV
def process_csv_file(path: Path, out_dir: Path) -> Optional[Path]:
    logger.info(f"Processando CSV: {path}")

    try:
        df = pd.read_csv(path, encoding="utf-8", low_memory=False)
    except:
        try:
            df = pd.read_csv(path, encoding="latin-1", low_memory=False)
        except Exception:
            logger.exception(f"Falha ao ler CSV {path}")
            return None

    stats = compute_descriptive_stats(df)
    summary = textual_summary(df)

    base = out_dir / f"csv_{path.stem}"
    safe_mkdir(base)

    stats.to_csv(base / f"{path.stem}_stats.csv", encoding="utf-8-sig")
    summary.to_csv(base / f"{path.stem}_summary.csv", index=False, encoding="utf-8-sig")

    return base

#DOCX
def extract_text_and_tables_docx(path: Path) -> Tuple[List[str], List[pd.DataFrame]]:
    if docx is None:
        raise RuntimeError("python-docx não instalado.")

    document = docx.Document(path)
    texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    tables = []
    for tb in document.tables:
        data = []
        for row in tb.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if len(data) > 1:
            df = pd.DataFrame(data[1:], columns=data[0])
        else:
            df = pd.DataFrame(data)
        tables.append(df)

    return texts, tables


def extract_numbers_from_text(texts: List[str]) -> List[float]:
    regex = re.compile(r"[-+]?\d[\d.,]*")
    nums = []
    for t in texts:
        for m in regex.findall(t):
            cleaned = m.replace(".", "").replace(",", ".")
            try:
                nums.append(float(cleaned))
            except:
                pass
    return nums


def process_docx_file(path: Path, out_dir: Path) -> Optional[Path]:
    logger.info(f"Processando DOCX: {path}")

    try:
        texts, tables = extract_text_and_tables_docx(path)
    except Exception:
        logger.exception(f"Erro ao processar DOCX {path}")
        return None

    base = out_dir / f"docx_{path.stem}"
    safe_mkdir(base)

    pd.DataFrame({"texto": texts}).to_csv(base / f"{path.stem}_text.csv", index=False, encoding="utf-8-sig")

    nums = extract_numbers_from_text(texts)
    if nums:
        pd.DataFrame(nums, columns=["valor"]).describe().to_csv(
            base / f"{path.stem}_numbers_stats.csv",
            encoding="utf-8-sig"
        )

    for i, df in enumerate(tables, start=1):
        df.to_csv(base / f"{path.stem}_table_{i}.csv", index=False, encoding="utf-8-sig")
        compute_descriptive_stats(df).to_csv(
            base / f"{path.stem}_table_{i}_stats.csv", encoding="utf-8-sig"
        )

    return base


#PDF
def extract_text_and_tables_pdf(path: Path) -> Tuple[str, List[pd.DataFrame]]:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber não instalado.")

    text_list = []
    tables = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt:
                text_list.append(txt)

            raw_tables = page.extract_tables()
            for t in raw_tables:
                if len(t) > 1:
                    tables.append(pd.DataFrame(t[1:], columns=t[0]))
                else:
                    tables.append(pd.DataFrame(t))

    return "\n".join(text_list), tables


def process_pdf_file(path: Path, out_dir: Path) -> Optional[Path]:
    logger.info(f"Processando PDF: {path}")

    try:
        text, tables = extract_text_and_tables_pdf(path)
    except Exception:
        logger.exception(f"Erro ao abrir PDF {path}")
        return None

    base = out_dir / f"pdf_{path.stem}"
    safe_mkdir(base)

    pd.DataFrame({"texto": [text]}).to_csv(base / f"{path.stem}_text.csv", index=False, encoding="utf-8-sig")

    for i, df in enumerate(tables, start=1):
        df.to_csv(base / f"{path.stem}_table_{i}.csv", encoding="utf-8-sig", index=False)
        stats = compute_descriptive_stats(df)
        stats.to_csv(base / f"{path.stem}_table_{i}_stats.csv", encoding="utf-8-sig")

    return base


#Excel
def process_excel_file(path: Path, out_dir: Path) -> Optional[Path]:
    logger.info(f"Processando Excel: {path}")

    try:
        xls = pd.ExcelFile(path)
    except:
        logger.exception(f"Falha ao abrir Excel {path}")
        return None

    base = out_dir / f"excel_{path.stem}"
    safe_mkdir(base)

    combined = []

    for sheet in xls.sheet_names:
        try:
            df = pd.read_excel(path, sheet_name=sheet)
            stats = compute_descriptive_stats(df)
            stats["sheet"] = sheet
            stats.to_csv(base / f"{path.stem}_{sheet}_stats.csv", encoding="utf-8-sig")
            combined.append(stats)
        except:
            logger.exception(f"Erro na aba {sheet} do arquivo {path}")

    if combined:
        pd.concat(combined).to_csv(base / f"{path.stem}_combined_stats.csv", encoding="utf-8-sig")

    return base


#Encontrar Arquivos
def find_files(directory: Path, types: List[str]) -> List[Path]:
    ext_map = {
        "csv": [".csv"],
        "docx": [".docx"],
        "pdf": [".pdf"],
        "xlsx": [".xlsx", ".xls", ".xlsm"]
    }

    files = []
    for t in types:
        for ext in ext_map.get(t, []):
            files.extend(directory.rglob(f"*{ext}"))

    return sorted(set(files))


#execução principal
def run_main(args):
    input_dir = args.input_dir
    out_dir = args.out_dir
    safe_mkdir(out_dir)

    logger.info(f"Input: {input_dir} | Output: {out_dir}")

    files = find_files(input_dir, args.types)
    if not files:
        logger.warning("Nenhum arquivo encontrado.")
        return

    handlers = {
        ".csv": process_csv_file,
        ".docx": process_docx_file,
        ".pdf": process_pdf_file,
        ".xlsx": process_excel_file,
        ".xls": process_excel_file,
        ".xlsm": process_excel_file,
    }

    for f in tqdm(files, desc="Processando arquivos"):
        handler = handlers.get(f.suffix.lower())
        if handler:
            try:
                handler(f, out_dir)
            except:
                logger.exception(f"Erro ao processar {f}")

    data_dir = Path("./data")
    copy_csvs_to_data(out_dir, data_dir)
    logger.info(f"CSVs também copiados para: {data_dir}")
    

    logger.info(f"Finalizado! Resultados em: {out_dir}")


parser = argparse.ArgumentParser(description="Processador de arquivos + Dashboard Streamlit")

parser.add_argument("--input-dir", "-i", type=Path, default=Path("./arquivos_documentos"))
parser.add_argument("--out-dir", "-o", type=Path, default=Path("./extraidos"))
parser.add_argument("--types", "-t", nargs="+", default=["csv", "docx", "pdf", "xlsx"])
parser.add_argument("--stream", action="store_true", help="Executar dashboard Streamlit")


if __name__ == "__main__":
    args = parser.parse_args()

    #Criar pasta de entrada se não existir
    if not args.input_dir.exists():
        args.input_dir.mkdir(parents=True, exist_ok=True)
        print("[AVISO] Pasta de entrada criada automaticamente. Coloque arquivos e rode novamente.")

    #Se pasta vazia
    if not any(p.is_file() for p in args.input_dir.rglob("*")):
        print("[AVISO] A pasta de entrada está vazia. Nenhum arquivo para processar.")

    else:
        run_main(args)
